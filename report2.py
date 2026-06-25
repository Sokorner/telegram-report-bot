
from telegram import Update
from telegram.ext import (
    ApplicationBuilder,
    CommandHandler,
    MessageHandler,
    ContextTypes,
    filters,
)

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

# =============================
# Store user data
# =============================
user_data = {}

# =============================
# Start command
# =============================
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "សូមស្វាគមន៍ 🙏\nវាយ /report ដើម្បីចាប់ផ្តើមរាយការណ៍"
    )

# =============================
# Report command
# =============================
async def report(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.message.from_user.id
    user_data[user_id] = {"step": "title"}

    # ✅ Instruction (Telegram only)
    await update.message.reply_text(
        "សូមបំពេញចំណងជើងរបាយការណ៍\n"
        "ឧទាហរណ៍៖ របាយការណ៍ស្តីពី… នៅថ្ងៃទី…"
    )

# =============================
# Handle messages
# =============================
async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.message.from_user.id
    text = update.message.text.strip()

    if user_id not in user_data:
        return

    step = user_data[user_id]["step"]

    # -------- TELEGRAM FLOW --------

    if step == "title":
        user_data[user_id]["title"] = text
        user_data[user_id]["step"] = "date"
        await update.message.reply_text(
            "កាលបរិច្ឆេទធ្វើរបាយការណ៍ (ចន្ទគតិ និងសុរិយាគតិ)"
        )

    elif step == "date":
        user_data[user_id]["date"] = text
        user_data[user_id]["step"] = "intro"
        await update.message.reply_text(
            "សូមបំពេញព័ត៌មានទូទៅ "
            "(ទីកន្លែង កាលបរិច្ឆេទ អធិបតីភាព អ្នកចូលរួម ។ល។)"
        )

    elif step == "intro":
        user_data[user_id]["intro"] = text
        user_data[user_id]["step"] = "agenda"
        await update.message.reply_text("សូមបំពេញរបៀបវារៈ")

    elif step == "agenda":
        user_data[user_id]["agenda"] = text
        user_data[user_id]["step"] = "result"
        await update.message.reply_text("ក្រោយពីពិភាក្សា តើទទួលបានលទ្ធផលអ្វីខ្លះ?")

    elif step == "result":
        user_data[user_id]["result"] = text
        user_data[user_id]["step"] = "note"
        await update.message.reply_text("សូមផ្តល់ការកត់សម្គាល់ ឬមតិរបស់អ្នក")

    elif step == "note":
        user_data[user_id]["note"] = text
        user_data[user_id]["step"] = "summary"
        await update.message.reply_text("សូមផ្តល់ការសន្និដ្ឋាន")

    elif step == "summary":
        user_data[user_id]["summary"] = text
        user_data[user_id]["step"] = "name"
        await update.message.reply_text("ឈ្មោះអ្នកធ្វើរបាយការណ៍")

    # -------- FINAL STEP: CREATE WORD --------

    elif step == "name":
        user_data[user_id]["name"] = text

        doc = Document()
        section = doc.sections[0]

        # Page setup (A4)
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(1.5)

        # ---------- Styles ----------
        def body(p):
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for r in p.runs:
                r.font.name = "Khmer OS Siemreap"
                r.font.size = Pt(12)

        def header(p):
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for r in p.runs:
                r.font.name = "Khmer OS Muol Light"
                r.font.size = Pt(12)

        # ---------- Header ----------
        p = doc.add_paragraph("ព្រះរាជាណាចក្រកម្ពុជា")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        header(p)

        p = doc.add_paragraph("ជាតិ សាសនា ព្រះមហាក្សត្រ")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        header(p)

        doc.add_paragraph("")

        header(doc.add_paragraph("ក្រសួងទេសចរណ៍"))
        header(doc.add_paragraph("អគ្គនាយកដ្ឋានរដ្ឋបាល និងហិរញ្ញវត្ថុ"))
        header(doc.add_paragraph("នាយកដ្ឋានសន្តិសុខ និងសុវត្ថិភាពទេសចរ"))

        # Date
        p = doc.add_paragraph(f"{user_data[user_id]['date']}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        body(p)

        # Title
        p = doc.add_paragraph(f"{user_data[user_id]['title']}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        body(p)

        # Content (NO instruction text)
        for k in ["intro", "agenda", "result", "note", "summary"]:
            body(doc.add_paragraph(user_data[user_id][k]))

        # Reporter
        p = doc.add_paragraph("អ្នកធ្វើរបាយការណ៍")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        body(p)

        p = doc.add_paragraph(user_data[user_id]["name"])
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        body(p)

        filename = f"{user_data[user_id]['name']}_report.docx"
        doc.save(filename)

        with open(filename, "rb") as f:
            await update.message.reply_document(document=f, filename=filename)

        user_data.pop(user_id)

# =============================
# Run bot
# =============================
app = ApplicationBuilder().token("8503397726:AAHHL2QsPPSmo3cevdNsSies1M80wBDRChw").build()

app.add_handler(CommandHandler("start", start))
app.add_handler(CommandHandler("report", report))
app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))

print("✅ Bot is running...")
app.run_polling()

